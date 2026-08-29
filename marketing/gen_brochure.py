#!/usr/bin/env python3
"""Mall Manager — Commercial Building Management System brochure (.docx)"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
SUB = RGBColor(0x2E, 0x74, 0xB5)
RED = RGBColor(0x7F, 0x1D, 0x1D)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs: r.font.color.rgb = ACCENT
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs: r.font.color.rgb = SUB
    return p

def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align: p.alignment = align
    return p

def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it)
        r.font.size = Pt(size)

def table(headers, rows, widths=None, size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(size)
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            c[i].text = str(v)
            for p in c[i].paragraphs:
                for r in p.runs: r.font.size = Pt(size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

def feature_group(title, intro, feats):
    h2(title)
    para(intro, italic=True, size=10.5, color=SUB)
    bullets(feats)

# ── COVER ──
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MALL MANAGER'); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RED
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Commercial Building Management System (BMS)'); r.bold = True; r.font.size = Pt(17); r.font.color.rgb = ACCENT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Shopping Mall · Commercial Complex · Market Society'); r.font.size = Pt(13); r.font.color.rgb = SUB
doc.add_paragraph()
para('One complete digital system for mall committees and property owners — automatic service-charge billing, sub-meter electricity/water billing, collections with instant receipts, full accounting, owner portal and Bangla-first interface.', italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('Prepared by: KRTaker (A concern of BITSCOL), Marketed By: Appvaley', bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ── MODULE MAP ──
h1('Mall Manager at a Glance')
doc.add_picture('/tmp/mall_shots/05_module_map.png', width=Inches(6.6))
doc.add_paragraph()
para('Thirteen modules — from meter readings to the annual general meeting — in one system.', italic=True, size=10.5, color=SUB)
doc.add_page_break()

# ── OVERVIEW ──
h1('1. Product Overview')
para('Mall Manager is a complete Commercial Building Management System (BMS) built for shopping malls, commercial complexes, and market societies. It replaces handwritten bill pads, paper ledgers and spreadsheet tracking with one integrated system that handles everything from meter readings to the annual general meeting — while printing bills that look exactly like the traditional committee bill pads shop owners already know.')
para('Designed for the local market: full বাংলা (Bangla) interface by default, ৳ currency handling, bKash/Nagad/bank/cash collections, DESCO/WASA utility billing, and printed receipts with committee signatures — all running on affordable shared hosting.')
table(['For', 'Key outcomes'],
      [['Mall & market owners’ societies / committees', 'Automatic monthly billing, higher collection rate, zero manual ledger'],
       ['Mall developers & managing agents', 'Multi-space portfolio, flexible ownership, full accounting visibility'],
       ['Commercial complex property managers', 'Sub-meter utility billing, vendor/AMC control, audit trail']],
      widths=[2.6, 3.9])

# ── MODULES ──
h1('2. Feature Modules')

feature_group('2.1 Space & Ownership Management',
  'Every shop, owner and tenant in one directory.',
  ['Commercial spaces with floor, size (sqft), billing model and service rate',
   'Flexible ownership — persons, companies, banks, trusts; one owner can hold many spaces (portfolio)',
   'Occupancy tracking: Active / Vacant / Closed with occupancy KPIs',
   'Tenant profiles + rental agreements (optional rent collection service)',
   'Space search by shop no / owner / mobile'])

feature_group('2.2 Service-Charge Billing Engine',
  'Automatic monthly bills with the exact charging model each committee uses.',
  ['Four billing models: Fixed (flat), Per sqft (rate × size), Fixed + utilities, Per sqft + utilities',
   'One-click monthly bill generation with due dates; manual sequence control',
   'Late-fee engine: %, grace days, minimum fine, cap — fully configurable',
   'Combined bill per space: service + electricity + water + fines on one printable form',
   'Printed bills replicate the traditional Bengali bill pad — বিদ্যুৎ/সার্ভিস চার্জ এবং অন্যান্য বিল — with numbered fields, dotted lines, rules and two committee signatures',
   'Print templates: A4, A5, and ½+½ A4 two-per-sheet (portrait/landscape), plus PDF download'])

feature_group('2.3 Sub-Meter Electricity & Water Billing',
  'Read meters, auto-generate utility bills, and know your true utility economics.',
  ['Monthly sub-meter readings with photo capture (photo is mandatory — spec compliance)',
   'Automatic units calculation and electricity/water bills at configurable ৳/unit rates',
   '200% usage anomaly flag for meter error / theft detection',
   'Effective electricity rate calculator: DESCO main bill ÷ total units → suggested rates (+2%/+5%/+10%)',
   'Utility own-income model: collections are income, DESCO/WASA bills are expenses — monthly profit/loss report'])

feature_group('2.4 Collections, Receipts & Payments',
  'Take money in any way your owners pay, and print a proper receipt on the spot.',
  ['Collect against bills (full/partial) with instant auto-numbered receipts (RCT-YYYYMM-####)',
   'Payment methods: cash, bank, bKash, Nagad — with MULTIPLE bank accounts (e.g. Brac Bank + EBL + bKash Business + Nagad)',
   'Money Receipt print: bank/A/C details, charge/late-fee lines, committee signatures',
   'Two-level approval workflow: waivers and payment voids are requested → approved by admin, receipt locked',
   'Rent collection (RNT-) and vendor payments (VNP-) with the same method/account precision'])

feature_group('2.5 Accounting — Complete Double-Entry Books',
  'A real chart of accounts, journal and reports — not just collections.',
  ['Multi-level Chart of Accounts (groups + sub-ledgers: Assets/Liabilities/Equity/Income/Expenses)',
   'Journal vouchers with double-entry posting; posting only to leaf accounts',
   'Trial Balance, P&L Statement, Party Ledger (owners/tenants/vendors/staff) with running balance',
   'Cashflow by month & method with per-account balances (per bank/mobile account)',
   'Smart Ledger auto-posts every collection, expense, salary and utility transaction',
   'Bank statement import (CSV) with automatic matching → reconciliation verdicts',
   'All reports printable/CSV'])

feature_group('2.6 Financial Views for the Committee',
  'Dashboards and invoice/payment views the committee actually checks.',
  ['Dashboard: collected vs outstanding, expenses, top defaulters, per-account cash balances',
   'Invoices view: combined per-space invoices (INV-YYYYMM-shop) with itemized service/elec/water/fines, print & PDF',
   'Payments view: every receipt (RCT + RNT) with payer, method, exact bank account, status (approved/pending void/voided)',
   'Analytics: billed vs collected trends, expense mix, occupancy, defaulters',
   'Period statements with opening/closing balance for any party'])

feature_group('2.7 Alerts, AMC & Risk Management',
  'The system watches the money — and the equipment.',
  ['High-dues auto alerts after configurable months (e.g. 2)',
   'Disconnection-risk detection (e.g. 3+ months overdue) with one-tap SMS warning',
   'AMC / servicing contract expiry tracking for lifts, generators, escalators, fire equipment — 30-day reminders',
   'Per-space due overview and Remind-all-defaulters SMS blast'])

feature_group('2.8 SMS & Communication',
  'Bangla SMS to owners and tenants — automatic and on demand.',
  ['Instant SMS receipt on every collection («মল»: রসিদ RCT-… — ৳… প্রাপ্ত হয়েছে)',
   'Dues reminders, disconnection warnings, notice broadcasts — all in Bangla',
   'Recipient control: owner, tenant, or both',
   'Provider-ready: bulksmsbd gateway (test mode available)'])

feature_group('2.9 Vendors, Expenses & Assets',
  'Everything the committee spends on — tracked to the ledger.',
  ['Vendor directory with categories, contracts and payment history (VNP- receipts)',
   'Expense entries with voucher photo upload, auto-posted to the ledger',
   'Assets register: install date, cost, warranty/AMC until',
   'Salary payments for staff/security guards with payroll view'])

feature_group('2.10 Complaints & Maintenance',
  'A proper ticket flow instead of phone calls.',
  ['Log complaints (lift / AC / light / water…) with priority',
   'Status workflow: Open → In Progress → Resolved; every change logged',
   'Space history: bills, readings, payments, complaints per shop'])

feature_group('2.11 Governance & Committee',
  'Formal records for the society.',
  ['Committee members & office bearers (chairman · secretary · treasurer)',
   'Meeting register with agenda, decisions and minutes',
   'Resolutions archived (spec 3.11)',
   'Notices with pin-to-top for announcements to all owners'])

feature_group('2.12 Security, Roles & Audit',
  'Right access, right approvals, full trail.',
  ['Roles: Super Admin · Owner · Manager · Accountant · Collector (limited to collections)',
   'Two-level approval for waivers and payment voids',
   'Full audit log: who did what, when',
   'License key system — vendor (Super Admin) controlled'])

feature_group('2.13 Owner Portal & Offline',
  'Self-service for owners, and it still works when the internet drops.',
  ['Owner portal: view own bills, dues, payments, notices; file complaints (Bangla)',
   'Offline + auto-sync: queued writes replay when back online; app-shell cache',
   'Works on mobile, tablet and desktop — responsive bottom-friendly UI'])

# ── FEATURE SUMMARY ──
h1('3. Feature Summary')
table(['Module', 'Highlights'],
      [['Billing', '4 billing models · auto generation · fine engine · combined bill print/PDF (A4/A5/½+½)'],
       ['Meters', 'Photo readings · units auto-calc · anomaly flag · effective-rate calculator · utility P/L'],
       ['Collections', 'Cash/bank/bKash/Nagad · multiple bank accounts · instant RCT- receipts · waivers & voids'],
       ['Accounting', 'Multi-level COA · double-entry journal · trial balance · P&L · party ledgers · cashflow · bank recon'],
       ['Reports', 'Dashboard · invoices · payments · analytics · statements · CSV/print/PDF'],
       ['Communication', 'Bangla SMS: receipts, reminders, warnings, blasts (bulksmsbd)'],
       ['Assets & AMC', 'Register, warranty, AMC expiry alerts'],
       ['Complaints', 'Ticketing with priority & status flow'],
       ['Governance', 'Committee · meetings · resolutions · notices'],
       ['Security', '5 roles · two-level approvals · audit trail · license keys'],
       ['Portal', 'Owner self-service (bills/dues/complaints)'],
       ['Localization', 'Bangla default + English toggle · ৳ formatting · Bangla printed forms'],
       ['Offline', 'Write queue + auto-sync · app-shell cache']],
      widths=[1.5, 5.0])

# ── BENEFITS ──
h1('4. Why Committees Choose Mall Manager')
bullets([
    'Collect MORE: automatic billing + overdue alerts + SMS reminders + disconnection warnings raise on-time collection dramatically',
    'Stop the ledger mess: every ৳ auto-posts to a real double-entry chart of accounts — trial balance always balanced',
    'Owners trust it: printed bills and receipts look like the traditional committee pad — familiar, formal, with signatures',
    'Bangla-first: the entire staff can use it in বাংলা; English mode for anyone else',
    'Cheap to run: single-file PHP + SQLite on ordinary shared hosting (cPanel) — no dedicated server, no database license',
    'Built for Bangladesh: bKash/Nagad/bank/cash, DESCO/WASA, BDT ৳, market-society governance',
    'Offline-safe: meter readings and collections never get lost during internet outages',
])

# ── SCREENSHOTS ──
h1('5. Product in Action')
para('Real screens from the live demo (রাজ্জাক প্লাজা) — Bangla interface.', italic=True, size=10.5, color=SUB)
for label, path in [
    ('Dashboard — collections, outstanding, expenses & cash balances', '/tmp/mall_shots/01_dashboard.png'),
    ('Invoices view — combined bills per space with print & PDF', '/tmp/mall_shots/02_invoices.png'),
    ('Payments view — every receipt with method, account & status', '/tmp/mall_shots/03_payments.png'),
    ('Printed bill — replicates the traditional committee bill pad', '/tmp/mall_shots/04_bill_form.png'),
]:
    h2(label)
    doc.add_picture(path, width=Inches(6.4))
    doc.add_paragraph()

# ── TECH ──
h1('6. Technical Profile')
table(['Aspect', 'Detail'],
      [['Architecture', 'Vue 3 single-page app + single-file PHP 8 + SQLite API'],
       ['Hosting', 'Standard shared hosting (cPanel / Apache / LiteSpeed) — no VPS required'],
       ['Devices', 'Mobile, tablet, desktop — responsive UI'],
       ['Printing', 'Browser print with A4 / A5 / ½+½ A4 templates + PDF export'],
       ['SMS', 'bulksmsbd gateway integration (Bangla templates)'],
       ['Offline', 'Service-worker app shell + localStorage write queue with auto-sync'],
       ['Language', 'বাংলা (default) / English'],
       ['Security', 'Bearer tokens, role-based access, approval workflows, audit log, license keys']],
      widths=[1.5, 5.0])

# ── LICENSING ──
h1('7. Licensing & Onboarding')
bullets([
    'Flexible licensing: one-off purchase, yearly subscription/license, or per-user monthly',
    'Vendor-controlled Super Admin: manage mall data, users and license from the dedicated control panel',
    'Onboarding: guided setup — mall profile, billing rules, account mapping, roles — then import shops & owners',
    'Demo data included so the committee can try every feature before going live',
])
doc.add_paragraph()
para('Mall Manager — Commercial Building Management System', bold=True, size=10, color=SUB, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Prepared by KRTaker (A concern of BITSCOL) · Marketed by Appvaley · mall.krtaker.com', size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save('/root/mall_brochure.docx')
print('saved /root/mall_brochure.docx')
